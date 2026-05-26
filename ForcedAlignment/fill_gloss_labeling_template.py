"""
fill_gloss_labeling_template.py
-------------------------------
Create deliverables that match Gloss_Labeling_Template.docx:

1. output/evaluation/evaluation_summary_template_format.csv
   Columns exactly match the template summary table.

2. Gloss_Labeling_Report_Filled.docx
   A filled copy of Gloss_Labeling_Template.docx with the ForcedAlignment
   full-run results.
"""
from __future__ import annotations

import argparse
import csv
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


FA_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = FA_DIR / "Gloss_Labeling_Template.docx"
DEFAULT_SUMMARY = FA_DIR / "output" / "evaluation" / "evaluation_summary.csv"
DEFAULT_OUT_DOCX = FA_DIR / "Gloss_Labeling_Report_Filled.docx"
DEFAULT_OUT_CSV = FA_DIR / "output" / "evaluation" / "evaluation_summary_template_format.csv"


EXPERIMENTS = {
    "1": {
        "title": "การทดลองที่ 1",
        "groundtruth": "CC_Aligned",
        "input": "CC",
        "output": "CC_Aligned_pred",
        "note": (
            "ใช้ CC tier และ whitespace tokenization; กรอง sil ออกจาก GT ในการประเมิน. "
            "Precision (67.5%) ต่ำกว่า Recall (69.8%) เล็กน้อยเพราะ CC บางคลิป "
            "(เช่น clip 982) ถูก whitespace.split() แตกออกหลาย tokens แต่ CC_Aligned "
            "มี 1 entry — predictions ส่วนเกินทำให้ตัวหารของ Precision ใหญ่ขึ้น"
        ),
    },
    "2": {
        "title": "การทดลองที่ 2",
        "groundtruth": "CC_Aligned",
        "input": "CC + sil",
        "output": "CC_Aligned_silmodel_pred",
        "note": (
            "CC_Aligned2 ไม่มีอยู่จริงใน dataset; ใช้ CC_Aligned พร้อมเพิ่ม sil หัว/ท้าย. "
            "ผลตกหนัก (F1 17.6% vs config 1 ที่ 68.6%) เพราะ architectural mismatch: "
            "SEA E4s-1 segmenter ตัด segment เฉพาะตอน active signing (ไม่สร้าง segment "
            "สำหรับ sil) → เมื่อเพิ่ม sil tokens ใน input, จำนวน tokens เกินจำนวน segments "
            "(K < T) ใน 85 clips → fallback uniform. รวม fallback clips จริง = 142 "
            "(57 จาก embedding หาย + 85 จาก K < T)"
        ),
    },
    "3": {
        "title": "การทดลองที่ 3",
        "groundtruth": "Gloss_Labeling",
        "input": "Gloss",
        "output": "Gloss_Labeling_pred",
        "note": (
            "Baseline สำหรับ Gloss tier; split ด้วย pipe delimiter. "
            "Mean IoU 0.2484 ดูต่ำเพราะ annotation convention: "
            "Gloss_Labeling ยืดทุก word ให้เต็มความยาวของ clip (avg GT duration = 3.84s) "
            "ส่วน aligner ทำนายช่วงเวลาที่สัญลักษณ์ปรากฏจริง (avg pred = 1.16s). "
            "Sanity check: text match 100%, 72.7% ของ predictions อยู่ภายใน GT, "
            "88.7% ของความยาว pred อยู่ใน GT — aligner หาตำแหน่งถูก แต่ GT กว้างกว่า"
        ),
    },
    "4": {
        "title": "การทดลองที่ 4",
        "groundtruth": "Gloss_Labeling1",
        "input": "Gloss1",
        "output": "Gloss_Labeling1_pred",
        "note": (
            "รวม sil token ใน input และ GT. Fallback clips จริง = 154 "
            "(57 จาก K=0 + 97 จาก K<T). "
            "F1 (20.7%) ดูดีกว่า config 3 (7.8%) แต่ DP จริง ๆ ทำได้แย่กว่า — "
            "split per-pair: DP-driven pairs mean IoU = 0.154 (hit 11.1%), "
            "fallback uniform pairs mean IoU = 0.709 (hit 88.3%). "
            "53% ของ hits ทั้งหมดมาจาก fallback uniform ที่บังเอิญตรง GT structure "
            "(sil+word+sil ใน GT ก็แบ่ง clip เท่า ๆ กันเหมือน uniform split)"
        ),
    },
    "5": {
        "title": "การทดลองที่ 5",
        "groundtruth": "Gloss_Labeling2",
        "input": "Gloss2",
        "output": "Gloss_Labeling2_pred",
        "note": (
            "รวม sil1/sil2 token ใน input และ GT. ผลเกือบเท่า config 4 "
            "(F1 20.9% vs 20.7%) — 91.7% ของ predictions identical กับ config 4. "
            "Numbered sil ไม่ช่วย DP เลือก segment ต่างกันมีนัยสำคัญ เพราะ SEA "
            "ไม่มี sil segments อยู่แล้ว. Fallback clips = 154 เท่ากับ config 4"
        ),
    },
    "6": {
        "title": "การทดลองที่ 6 - แยกทดสอบทีละประโยค",
        "groundtruth": "Gloss_Labeling",
        "input": "Gloss",
        "output": "Gloss_Labeling_pred",
        "note": "ไม่ได้รันแยก: dataset นี้ 1 clip = 1 sentence จึงเทียบเท่า experiment #3",
    },
}


def pct(value: str | float, digits: int = 1) -> str:
    return f"{float(value) * 100:.{digits}f}%"


def dec(value: str | float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def load_summary(path: Path) -> dict[str, dict]:
    with open(path, encoding="utf-8-sig", newline="") as f:
        return {row["config"]: row for row in csv.DictReader(f)}


def template_rows(summary: dict[str, dict]) -> list[dict]:
    rows: list[dict] = []
    for key in ["1", "2", "3", "4", "5"]:
        item = EXPERIMENTS[key]
        result = summary[key]
        rows.append({
            "การทดลอง": f"#{key}",
            "GroundTruth": item["groundtruth"],
            "Input": item["input"],
            "Precision": pct(result["precision_iou_0_5"]),
            "Recall": pct(result["recall_iou_0_5"]),
            "F1-Score": pct(result["f1_iou_0_5"]),
            "Accuracy": pct(result["accuracy_iou_0_5"]),
            "Mean IoU": dec(result["mean_iou"]),
            "Frame Accuracy": pct(result["frame_accuracy"]),
            "Fallback": result["fallback_predictions"],
            "หมายเหตุ": item["note"],
        })

    exp3 = summary["3"]
    item = EXPERIMENTS["6"]
    rows.append({
        "การทดลอง": "#6",
        "GroundTruth": item["groundtruth"],
        "Input": item["input"],
        "Precision": f"{pct(exp3['precision_iou_0_5'])}*",
        "Recall": f"{pct(exp3['recall_iou_0_5'])}*",
        "F1-Score": f"{pct(exp3['f1_iou_0_5'])}*",
        "Accuracy": f"{pct(exp3['accuracy_iou_0_5'])}*",
        "Mean IoU": f"{dec(exp3['mean_iou'])}*",
        "Frame Accuracy": f"{pct(exp3['frame_accuracy'])}*",
        "Fallback": f"{exp3['fallback_predictions']}*",
        "หมายเหตุ": item["note"],
    })
    return rows


def write_template_csv(rows: list[dict], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = [
        "การทดลอง", "GroundTruth", "Input", "Precision", "Recall",
        "F1-Score", "Accuracy", "Mean IoU", "Frame Accuracy",
        "Fallback", "หมายเหตุ",
    ]
    with open(out_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def set_cell(cell, text: str, bold: bool = False, align_center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(10)


def fill_overview(doc: Document) -> None:
    table = doc.tables[0]
    values = {
        1: "SEA SignCLIP multilingual + SEA E4s-1 segmentation + monotonic DP",
        2: "ForcedAlignment dataset",
        3: "1,132 clips; 1,132 EAF files; 1,132 MP4 files",
        4: "Positional IoU-only matching at threshold 0.5; report Precision, Recall, F1-Score, Accuracy",
        5: "| (pipe character) for Gloss tiers; whitespace for CC tier",
    }
    for row_idx, text in values.items():
        set_cell(table.cell(row_idx, 1), text)
    style_table(table)


def fill_experiment_tables(doc: Document, rows: list[dict]) -> None:
    # Tables 1-4 are experiments 1-4. Table 5 contains experiments 5 and 6.
    table_map = {"1": doc.tables[1], "2": doc.tables[2], "3": doc.tables[3], "4": doc.tables[4]}
    row_by_key = {row["การทดลอง"].lstrip("#"): row for row in rows}

    for key, table in table_map.items():
        item = EXPERIMENTS[key]
        result = row_by_key[key]
        set_cell(table.cell(1, 1), item["groundtruth"], align_center=True)
        set_cell(table.cell(1, 3), item["input"], align_center=True)
        set_cell(table.cell(2, 1), item["output"], align_center=True)
        note = (
            f"{item['note']}\n"
            f"P/R/F1/Acc={result['Precision']}/{result['Recall']}/{result['F1-Score']}/{result['Accuracy']}; "
            f"mIoU={result['Mean IoU']}; fallback={result['Fallback']}"
        )
        set_cell(table.cell(3, 1), note)
        style_table(table)

    # Template table 5 has a few merged cells, so use row.cells instead of
    # table.cell(row, col), which can be unstable on the final merged row.
    table = doc.tables[5]

    item = EXPERIMENTS["5"]
    result = row_by_key["5"]
    set_cell(table.rows[1].cells[1], item["groundtruth"], align_center=True)
    set_cell(table.rows[1].cells[3], item["input"], align_center=True)
    set_cell(table.rows[2].cells[1], item["output"], align_center=True)
    note = (
        f"{item['note']}\n"
        f"P/R/F1/Acc={result['Precision']}/{result['Recall']}/{result['F1-Score']}/{result['Accuracy']}; "
        f"mIoU={result['Mean IoU']}; fallback={result['Fallback']}"
    )
    set_cell(table.rows[3].cells[1], note)

    item = EXPERIMENTS["6"]
    result = row_by_key["6"]
    set_cell(table.rows[5].cells[1], item["groundtruth"], align_center=True)
    set_cell(table.rows[5].cells[3], item["input"], align_center=True)
    set_cell(table.rows[6].cells[1], item["output"], align_center=True)
    note = (
        f"{item['note']}\n"
        f"P/R/F1/Acc={result['Precision']}/{result['Recall']}/{result['F1-Score']}/{result['Accuracy']}; "
        f"mIoU={result['Mean IoU']}; fallback={result['Fallback']}"
    )
    set_cell(table.rows[7].cells[1], note)
    style_table(table)


def fill_summary_table(doc: Document, rows: list[dict]) -> None:
    table = doc.tables[6]
    for idx, row in enumerate(rows, start=1):
        set_cell(table.cell(idx, 0), row["การทดลอง"], align_center=True)
        set_cell(table.cell(idx, 1), row["GroundTruth"], align_center=True)
        set_cell(table.cell(idx, 2), row["Input"], align_center=True)
        set_cell(table.cell(idx, 3), row["Precision"], align_center=True)
        set_cell(table.cell(idx, 4), row["Recall"], align_center=True)
        set_cell(table.cell(idx, 5), row["F1-Score"], align_center=True)
        set_cell(table.cell(idx, 6), row["Accuracy"], align_center=True)
    style_table(table)


def fill_conclusion(doc: Document) -> None:
    table = doc.tables[7]
    set_cell(table.cell(0, 1), "การทดลองที่ 1 (CC -> CC_Aligned, no sil)")
    set_cell(
        table.cell(1, 1),
        "ได้ผลดีที่สุดใน full dataset: F1@0.5 = 68.6%, Accuracy = 69.8%, "
        "Mean IoU = 0.5928, Frame Accuracy = 76.8%. "
        "เป็น config ที่ความยาวของ pred (1.7s) และ GT (1.9s) ใกล้เคียงกันที่สุด "
        "ทำให้ IoU วัดได้แม่นยำ",
    )
    set_cell(
        table.cell(2, 1),
        "1. Config #1 ชนะทุก metric หลักใน dataset นี้ — DP ทำงานได้ดี (DP-driven mean IoU 0.608, hit 73.5%)\n"
        "2. ทุก config ที่มี sil ใน input (#2, #4, #5) มี K<T fallback: "
        "Distinct fallback clips: #1=57, #2=142 (+85), #3=60, #4=154 (+97), #5=154 (+97). "
        "K=0 ใน 57 clips (SEA ไม่ detect sign เลย), K=1 ใน 40 clips, K=2 ใน 45 clips → "
        "Config 2 (T=3) ตกหล่น 142 clips จาก K<3\n"
        "3. ที่สำคัญกว่า: DP-driven pairs ทำได้แย่ในทุก sil-config — "
        "Config 2 DP mean IoU = 0.127 (hit 7%), Config 4 DP mean IoU = 0.154 (hit 11.1%) "
        "เทียบกับ Config 1 DP = 0.608 (hit 73.5%). "
        "sil tokens ทำให้ DP สับสน (per-token softmax บังคับให้จับ segment แม้ไม่มี match จริง)\n"
        "4. F1 ของ Config #2/#4/#5 ที่ดูสูงกว่า 'DP-only baseline' จริง ๆ มาจาก "
        "fallback uniform ที่บังเอิญตรงโครงสร้าง GT — Config 4 fallback hit 88.3%, "
        "Config 2 fallback hit 92.5% (uniform split ทำให้ pred = GT structure เพราะ "
        "ทั้งคู่แบ่ง clip เท่า ๆ กันด้วย sil tokens)\n"
        "5. Config #3 (Gloss, no sil) Mean IoU 0.2484 ดูต่ำกว่า 04.mp4 baseline (0.4901) "
        "แต่ไม่ใช่ aligner regression — text match 100%, 72.7% ของ predictions อยู่ใน GT span, "
        "88.7% ของความยาว pred อยู่ใน GT. สาเหตุคือ Gloss_Labeling ยืดทุก word ให้ครอบคลุม "
        "ช่วงเวลายาว (avg 3.84s) ขณะที่ aligner ทำนายช่วงสัญลักษณ์จริง (avg 1.16s)",
    )
    set_cell(
        table.cell(3, 1),
        "1. ใช้ Config #1 เป็นผลหลักของรายงาน (F1 68.6%, mIoU 0.5928) — เทียบได้กับ 04.mp4 "
        "baseline (0.4901) → ดีกว่า. DP-driven hit rate 73.5% สื่อว่า aligner แม่นยำจริงเมื่อ "
        "input ไม่มี sil tokens รบกวน\n"
        "2. อย่ารายงาน Config #2/#4/#5 เป็น 'sil modeling แย่' — bottleneck คือ DP ที่สับสน "
        "เมื่อมี sil tokens ใน input ไม่ใช่เพราะ sil concept เอง. ทางแก้:\n"
        "   (A) Post-processing: ใช้ Config #1 หรือ #3 ทำ word alignment ก่อน แล้วเติม "
        "sil intervals ระหว่าง word predictions ในขั้นตอนแยก (ไม่ผ่าน DP เลย)\n"
        "   (B) แก้ DP cost: ให้ sil tokens skip embedding หรือใช้ uniform similarity "
        "เพื่อไม่บังคับ DP เลือก word-segment ผิด\n"
        "   (C) Fine-tune SEA E4s-1 ให้ output multi-class (sign/sil/transition)\n"
        "3. ระมัดระวังเมื่อเทียบ Config #3 mIoU 0.2484 กับ 04.mp4 0.4901 โดยตรง — "
        "ตัวเลขต่ำเพราะ annotation convention ต่างกัน (ForcedAlignment ยืด GT ให้เต็ม clip, "
        "04.mp4 ใช้ช่วงสัญลักษณ์จริง) ไม่ใช่ aligner regression\n"
        "4. สำหรับ Gloss config แนะนำ Frame Accuracy (26.2%) หรือ % any-overlap (97.3%) "
        "แทน IoU@0.5 เพราะไม่อ่อนไหวต่อความกว้างของ GT\n"
        "5. 57 clips ที่ SEA detect K=0 segments เป็น hard limit ของ E4s-1 — "
        "ปรับ threshold sign-b=30/sign-o=50 หรือใช้ pose-based heuristic fallback "
        "(เช่น hand motion peaks) แทน uniform split",
    )
    style_table(table)


def add_metric_note(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("หมายเหตุการกรอกผล: ").bold = True
    p.add_run(
        "ค่าที่กรอกในตารางสรุปใช้ IoU-only positional matching ที่ threshold 0.5 "
        "(ไม่ใช้ text-match เพื่อหลีกเลี่ยง structural leakage). "
        "Experiment #6 ทำเครื่องหมาย * เพราะไม่ได้ rerun แยก และถือว่าเทียบเท่า #3 บน dataset นี้."
    )


def keep_conclusion_together(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.strip().startswith("4. การวิเคราะห์"):
            para.paragraph_format.page_break_before = True
            break


def fill_docx(template_path: Path, out_path: Path, rows: list[dict]) -> None:
    shutil.copy2(template_path, out_path)
    doc = Document(out_path)
    fill_overview(doc)
    fill_experiment_tables(doc, rows)
    fill_summary_table(doc, rows)
    fill_conclusion(doc)
    keep_conclusion_together(doc)
    add_metric_note(doc)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill Gloss_Labeling_Template.docx from FA evaluation results.")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--summary", type=Path, default=DEFAULT_SUMMARY)
    parser.add_argument("--out-docx", type=Path, default=DEFAULT_OUT_DOCX)
    parser.add_argument("--out-csv", type=Path, default=DEFAULT_OUT_CSV)
    args = parser.parse_args()

    summary = load_summary(args.summary)
    rows = template_rows(summary)
    write_template_csv(rows, args.out_csv)
    fill_docx(args.template, args.out_docx, rows)

    print(f"[OK] Template CSV -> {args.out_csv}")
    print(f"[OK] Filled DOCX  -> {args.out_docx}")
    print(f"[OK] Completed at {datetime.now().isoformat(timespec='seconds')}")


if __name__ == "__main__":
    main()
